from __future__ import annotations

from io import BytesIO

from docx import Document

from app.services.input.exceptions import InputLoadError


def load_docx_bytes(content: bytes) -> str:
    if not content:
        raise InputLoadError("DOCX file is empty.")
    try:
        doc = Document(BytesIO(content))
    except Exception as exc:
        raise InputLoadError("Could not read DOCX document.") from exc
    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise InputLoadError("DOCX contains no extractable text.")
    return text
